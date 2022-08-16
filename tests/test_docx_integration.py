import docx
from pyommlbuilder.main import (
    MathPara,
    Math,
    Run,
    Text,
    Fraction,
    Numerator,
    Denominator,
    SquareRoot,
    SuperscriptObject,
    wrap_text,
)
from pyommlbuilder.helpers import hash_file, make_aligned_equation


def test_simple_expresson():
    expression = Math(Run(Text("x+3=10")))
    xml_element = expression._as_xml_element()

    doc = docx.Document()
    p = doc.add_paragraph()
    p._element.append(xml_element)
    doc_blob = doc._part.blob
    assert len(doc_blob) == 1606
    assert b"w:r" in doc_blob
    assert b"x+3" in doc_blob


def test_quadratic_equation():
    expression = Math(
        "x=",
        Fraction(
            Numerator(
                "-b±",
                SquareRoot(
                    [
                        SuperscriptObject(
                            "b",
                            "2",
                        ),
                        "-4ac",
                    ]
                ),
            ),
            Denominator("2a"),
        ),
    )

    xml_element = expression._as_xml_element()

    doc = docx.Document()
    p = doc.add_paragraph()
    p._element.append(xml_element)
    doc_blob = doc._part.blob
    assert len(doc_blob) == 1883
    assert b"degHide" in doc_blob
    assert b"m:num" in doc_blob
    assert b"4ac" in doc_blob
    assert b"\xc2\xb1" in doc_blob


def test_make_aligned_equation_block():
    line1 = make_aligned_equation("x+3", 8)
    line2 = make_aligned_equation("x", "8-3")
    line3 = make_aligned_equation("", 5)

    equation_block = MathPara(line1, line2, line3)

    xml_element = equation_block._as_xml_element()

    doc = docx.Document()
    p = doc.add_paragraph()
    p._element.append(xml_element)
    doc_blob = doc._part.blob
    print(doc_blob)
    assert len(doc_blob) == 1938
    assert b"<m:aln/>" in doc_blob
    assert b"<w:br/>" in doc_blob
